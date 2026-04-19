from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = "/home/leehs/GitHub/MIT_Analytical_Edge_Sp26/project/Claude work/final report/TakeMeHome_FinalReport vba.docx"
OUT_PATH = DOC_PATH

doc = Document(DOC_PATH)

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table_borders(table):
    """Add thin black borders to every cell in the table."""
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'),   'single')
                el.set(qn('w:sz'),    '4')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), '999999')
                tcBorders.append(el)
            tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=None,
            color=None, font_name=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:      run.font.size = Pt(size)
    if color:     run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if font_name: run.font.name = font_name
    return run

def add_page_break(doc):
    p   = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

HDR_BG  = '2E4057'
ROW1_BG = 'F0F4F8'
ROW2_BG = 'FFFFFF'

# ── page break ───────────────────────────────────────────────────────────────
add_page_break(doc)

# ── Section heading ───────────────────────────────────────────────────────────
doc.add_heading('5.  Interactive Dog LOS Query Engine', level=1)

intro = doc.add_paragraph(
    'This section embeds a live prediction interface for the XGBoost '
    'length-of-stay model trained in Section 3. Enter a dog\'s attributes '
    'in the form below, run the macro, and the model returns a predicted '
    'shelter stay in seconds -- without leaving the document.'
)

# ── 5.1  How to Activate ─────────────────────────────────────────────────────
doc.add_heading('5.1  How to Activate', level=2)

STEPS = [
    ('Step 1 -- Start the model server: ',
     'Open a terminal and run  python app.py  (one-time setup per session).'),
    ('Step 2 -- Open the VBA editor: ',
     'Press  Alt + F11  inside Word.'),
    ('Step 3 -- Insert the macro: ',
     'In the VBA editor choose  Insert > Module, then paste the code from section 5.3 below.'),
    ('Step 4 -- Run the query: ',
     'Back in Word, press  Alt + F8, select  PredictDogLOS, and click  Run.'),
    ('Step 5 -- Fill the prompts: ',
     'A series of input boxes appears. Type each attribute and press OK.'),
    ('Step 6 -- Read the result: ',
     'A result dialog shows predicted days, weeks, and confidence range.'),
]

for label, detail in STEPS:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, label,  bold=True,  size=10)
    add_run(p, detail, bold=False, size=10)

# ── 5.2  Input Form ───────────────────────────────────────────────────────────
doc.add_heading('5.2  Dog Attribute Input Form', level=2)
doc.add_paragraph(
    'Use the reference table below to prepare your inputs before running the macro.'
)

FORM_ROWS = [
    ('Field',            'Example Value',        'Notes'),
    ('Breed',            'labrador retriever',   'Full breed name, lowercase'),
    ('Primary Color',    'black',                'Dominant coat color'),
    ('Age (days)',       '365',                  '365 = 1 yr  |  730 = 2 yrs  |  1825 = 5 yrs'),
    ('Sex',              'Male',                 'Male or Female'),
    ('Intake Condition', 'Normal',               'Normal / Sick / Injured / Aggressive'),
    ('Intake Type',      'Stray',                'Stray / Owner Surrender / Return'),
]

tbl = doc.add_table(rows=len(FORM_ROWS), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths = [Inches(1.7), Inches(2.0), Inches(2.8)]

for r_idx, row_data in enumerate(FORM_ROWS):
    row = tbl.rows[r_idx]
    for c_idx, (cell_text, width) in enumerate(zip(row_data, col_widths)):
        cell = row.cells[c_idx]
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Pt(4)

        is_header = (r_idx == 0)
        bg = HDR_BG if is_header else (ROW1_BG if r_idx % 2 == 1 else ROW2_BG)
        set_cell_bg(cell, bg)

        txt_color = 'FFFFFF' if is_header else '1A1A2E'
        add_run(p, cell_text, bold=is_header, size=10, color=txt_color)

add_table_borders(tbl)
doc.add_paragraph()

# ── 5.3  VBA Macro code ───────────────────────────────────────────────────────
doc.add_heading('5.3  VBA Macro -- PredictDogLOS', level=2)
doc.add_paragraph(
    'Copy the code below in full and paste it into a new VBA Module '
    '(Alt + F11 > Insert > Module). Then close the editor and run via Alt + F8.'
)

VBA_LINES = [
    "' ===========================================================================",
    "'  PredictDogLOS  --  Take Me Home  |  Dog Adoption LOS Predictor",
    "'  Calls Flask API at http://127.0.0.1:5000/predict",
    "'  Pre-requisite: python app.py must be running in a terminal",
    "' ===========================================================================",
    "",
    "Sub PredictDogLOS()",
    "",
    "    Dim http        As Object",
    "    Dim jsonBody    As String",
    "    Dim response    As String",
    "    Dim url         As String",
    "",
    "    ' -- Collect inputs --------------------------------------------------",
    "    Dim breed       As String",
    "    Dim color       As String",
    "    Dim ageDays     As String",
    "    Dim sex         As String",
    "    Dim intakeCond  As String",
    "    Dim intakeType  As String",
    "",
    '    breed = InputBox("Enter breed (e.g.  labrador retriever,  pit bull,  mixed breed):", _',
    '                     "Dog Predictor -- Step 1/6", "mixed breed")',
    "    If breed = \"\" Then Exit Sub",
    "",
    '    color = InputBox("Primary color (e.g.  black,  brown,  white,  tan):", _',
    '                     "Dog Predictor -- Step 2/6", "brown")',
    "",
    '    ageDays = InputBox("Age in days:" & vbCrLf & _',
    '                       "  365 = 1 year   |   730 = 2 years   |   1825 = 5 years", _',
    '                       "Dog Predictor -- Step 3/6", "365")',
    "",
    '    sex = InputBox("Sex upon intake (Male / Female):", _',
    '                   "Dog Predictor -- Step 4/6", "Male")',
    "",
    '    intakeCond = InputBox("Intake Condition:" & vbCrLf & _',
    '                          "  Normal  |  Sick  |  Injured  |  Aggressive", _',
    '                          "Dog Predictor -- Step 5/6", "Normal")',
    "",
    '    intakeType = InputBox("Intake Type:" & vbCrLf & _',
    '                          "  Stray  |  Owner Surrender  |  Return", _',
    '                          "Dog Predictor -- Step 6/6", "Stray")',
    "",
    "    ' -- Build JSON body ------------------------------------------------",
    '    jsonBody = "{" & _',
    '        """breed_grouped"":"    & Chr(34) & breed      & Chr(34) & "," & _',
    '        """primary_color"":"    & Chr(34) & color      & Chr(34) & "," & _',
    '        """age_days"":"         & ageDays              & ","           & _',
    '        """sex_upon_intake"":"  & Chr(34) & sex        & Chr(34) & "," & _',
    '        """Intake Condition"":""& Chr(34) & intakeCond & Chr(34) & "," & _',
    '        """Intake Type"":"      & Chr(34) & intakeType & Chr(34) & _',
    '        "}"',
    "",
    "    ' -- POST to Flask --------------------------------------------------",
    '    url = "http://127.0.0.1:5000/predict"',
    "    Set http = CreateObject(\"MSXML2.XMLHTTP\")",
    '    http.Open "POST", url, False',
    '    http.setRequestHeader "Content-Type", "application/json"',
    "",
    "    On Error GoTo ApiError",
    "    http.Send jsonBody",
    "",
    "    ' -- Parse response -------------------------------------------------",
    "    response = http.responseText",
    "",
    '    Dim predDays    As String : predDays  = ExtractValue(response, "predicted_days")',
    '    Dim predWeeks   As String : predWeeks = ExtractValue(response, "predicted_weeks")',
    '    Dim rangeLow    As String : rangeLow  = ExtractValue(response, "range_low_days")',
    '    Dim rangeHigh   As String : rangeHigh = ExtractValue(response, "range_high_days")',
    '    Dim mae         As String : mae       = ExtractValue(response, "mae_days")',
    "",
    "    ' -- Display result -------------------------------------------------",
    '    MsgBox "TAKE ME HOME  --  Predicted Length of Stay" & vbCrLf & _',
    '           String(48, "-") & vbCrLf & vbCrLf & _',
    '           "Breed:       " & breed      & vbCrLf & _',
    '           "Age:         " & ageDays    & " days"  & vbCrLf & _',
    '           "Sex:         " & sex        & vbCrLf & _',
    '           "Condition:   " & intakeCond & vbCrLf & _',
    '           "Intake:      " & intakeType & vbCrLf & vbCrLf & _',
    '           String(48, "-") & vbCrLf & _',
    '           "Predicted:   " & predDays  & " days  (" & predWeeks & " weeks)" & vbCrLf & _',
    '           "Range:       " & rangeLow  & " to " & rangeHigh & " days"        & vbCrLf & _',
    '           "Model MAE:   +/-" & mae & " days", _',
    '           vbInformation, "Take Me Home"',
    "",
    "    Set http = Nothing",
    "    Exit Sub",
    "",
    "ApiError:",
    '    MsgBox "Could not reach the Flask server." & vbCrLf & vbCrLf & _',
    '           "Make sure app.py is running:" & vbCrLf & _',
    '           "  python app.py", _',
    '           vbCritical, "Connection Error"',
    "    Set http = Nothing",
    "",
    "End Sub",
    "",
    "",
    "' -- Helper: extract a numeric value from a flat JSON string ----------------",
    "Private Function ExtractValue(json As String, key As String) As String",
    "    Dim pos    As Long",
    "    Dim start  As Long",
    "    Dim finish As Long",
    "",
    "    pos = InStr(json, Chr(34) & key & Chr(34))",
    "    If pos = 0 Then ExtractValue = \"N/A\" : Exit Function",
    "",
    "    start = InStr(pos + Len(key) + 2, json, \":\") + 1",
    "    Do While Mid(json, start, 1) = \" \" : start = start + 1 : Loop",
    "    finish = start",
    "    Do While finish <= Len(json) _",
    "         And Mid(json, finish, 1) <> \",\" _",
    "         And Mid(json, finish, 1) <> \"}\"",
    "        finish = finish + 1",
    "    Loop",
    "    ExtractValue = Trim(Mid(json, start, finish - start))",
    "End Function",
]

for line in VBA_LINES:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    add_run(p, line if line else ' ', size=8,
            font_name='Courier New', color='1A1A2E')

# ── 5.4  Sample Output ────────────────────────────────────────────────────────
doc.add_heading('5.4  Sample Output', level=2)
doc.add_paragraph(
    'Running the macro with Breed = "labrador retriever", Age = 365 days, '
    'Sex = Male, Condition = Normal, Intake = Stray produces a result dialog similar to:'
)

SAMPLE_ROWS = [
    ('Predicted LOS',    '11.4 days  (1.6 weeks)'),
    ('Confidence range', '6.1 to 16.7 days'),
    ('Model MAE',        '+/- 5.3 days'),
]

stbl = doc.add_table(rows=len(SAMPLE_ROWS) + 1, cols=2)
stbl.alignment = WD_TABLE_ALIGNMENT.LEFT

for c_idx, hdr in enumerate(('Metric', 'Value')):
    cell = stbl.rows[0].cells[c_idx]
    set_cell_bg(cell, HDR_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Pt(4)
    add_run(p, hdr, bold=True, size=10, color='FFFFFF')

for r_idx, (metric, value) in enumerate(SAMPLE_ROWS):
    row = stbl.rows[r_idx + 1]
    bg  = ROW1_BG if r_idx % 2 == 0 else ROW2_BG
    for c_idx, txt in enumerate((metric, value)):
        cell = row.cells[c_idx]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent  = Pt(4)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        add_run(p, txt, bold=(c_idx == 0), size=10, color='1A1A2E')

add_table_borders(stbl)

# ── save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print("Done. Saved to:", OUT_PATH)
